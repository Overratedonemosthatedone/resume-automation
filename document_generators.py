"""
Document Generation Module

Converts plain text resumes into .docx and PDF formats.
Maintains clean, ATS-friendly formatting in both outputs.
"""

import json
import re
from datetime import datetime
from pathlib import Path

import config
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from loguru import logger


# ============================================================================
# DOCX Generator
# ============================================================================

class DocxGenerator:
    """Generate .docx (Microsoft Word) resume files"""
    
    def __init__(self, template_path=None):
        """
        Initialize generator.
        
        Args:
            template_path: Optional path to template .docx file
        """
        if template_path and Path(template_path).exists():
            self.doc = Document(template_path)
            logger.info(f"Loaded template: {template_path}")
        else:
            self.doc = Document()
            self._set_default_styles()
    
    def _set_default_styles(self):
        """Set clean, professional, ATS-friendly default styles."""
        # Normal text style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0, 0, 0)
        
        # Heading style for section headers
        try:
            heading = self.doc.styles['Heading 1']
            heading.font.name = 'Calibri'
            heading.font.size = Pt(12)
            heading.font.bold = True
            heading.font.color.rgb = RGBColor(0, 0, 0)
        except:
            pass
    
    def add_text(self, text, style=None, bold=False, size=None):
        """
        Add a paragraph of text.
        
        Args:
            text: Text to add
            style: Paragraph style name
            bold: Whether to bold the text
            size: Font size in points
        """
        p = self.doc.add_paragraph(text, style=style)
        
        if bold or size:
            for run in p.runs:
                if bold:
                    run.bold = True
                if size:
                    run.font.size = Pt(size)
    
    def add_section_header(self, header_text):
        """Add a section header (e.g., 'EXPERIENCE', 'EDUCATION')."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        
        run = p.add_run(header_text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    def add_bullet(self, text, indent_level=0):
        """Add a bullet point."""
        p = self.doc.add_paragraph(text, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25 * (1 + indent_level))
        p.paragraph_format.space_after = Pt(3)
    
    def add_job_entry(self, title, company, location, dates, achievements):
        """
        Add a job entry with title, company, location, dates, and bullet points.
        
        Args:
            title: Job title
            company: Company name
            location: Location
            dates: Date range (e.g., "Jan 2022 - Present")
            achievements: List of achievement strings (will be bullet points)
        """
        # Title and company
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        
        title_run = p.add_run(title)
        title_run.font.bold = True
        title_run.font.size = Pt(11)
        
        p.add_run(f"\n{company}")
        
        # Location and dates
        p = self.doc.add_paragraph(f"{location} | {dates}")
        p.paragraph_format.space_after = Pt(3)
        
        # Achievements as bullets
        for achievement in achievements:
            self.add_bullet(achievement)
    
    def add_education_entry(self, degree, institution, location, year, gpa=None):
        """Add an education entry."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        
        degree_run = p.add_run(degree)
        degree_run.font.bold = True
        
        p.add_run(f"\n{institution}, {location}")
        
        gpa_text = f" | GPA: {gpa}" if gpa else ""
        p.add_run(f"\nGraduated: {year}{gpa_text}")
        
        p.paragraph_format.space_after = Pt(3)
    
    def parse_and_add_resume(self, resume_text):
        """
        Parse plain text resume and add to document.
        
        Expects format:
        - ALL CAPS lines = section headers
        - Lines starting with spaces = bullet points
        - Other lines = regular text
        """
        lines = resume_text.strip().split('\n')
        
        for line in lines:
            line = line.rstrip()
            
            if not line.strip():
                # Blank line - add space
                self.doc.add_paragraph()
            
            elif line.isupper() or line.endswith(':'):
                # Section header
                self.add_section_header(line)
            
            elif line.startswith('  ') or line.startswith('\t'):
                # Bullet point
                clean_line = line.lstrip()
                if clean_line.startswith('- ') or clean_line.startswith('• '):
                    clean_line = clean_line[2:].strip()
                self.add_bullet(clean_line)
            
            else:
                # Regular text
                self.add_text(line)
    
    def save(self, file_path):
        """
        Save document to .docx file.
        
        Args:
            file_path: Output file path
        """
        file_path = Path(file_path)
        file_path.parent.mkdir(parents=True, exist_ok=True)
        
        self.doc.save(str(file_path))
        logger.info(f"Saved DOCX: {file_path}")
        print(f"Saved resume to: {file_path}")


# ============================================================================
# PDF Generator
# ============================================================================

class PdfGenerator:
    """Generate PDF resume files with clean formatting"""
    
    def __init__(self):
        """Initialize PDF generator with custom styles."""
        self.styles = getSampleStyleSheet()
        self._add_custom_styles()
    
    def _add_custom_styles(self):
        """Add ATS-friendly styles for PDF output."""
        # Normal paragraph style
        self.styles.add(ParagraphStyle(
            name='ResumeNormal',
            fontName='Helvetica',
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
            spaceAfter=3,
        ))
        
        # Section header style
        self.styles.add(ParagraphStyle(
            name='ResumeSectionHeader',
            fontName='Helvetica-Bold',
            fontSize=11,
            leading=12,
            alignment=TA_LEFT,
            spaceAfter=3,
            spaceBefore=6,
            textColor=RGBColor(0, 0, 0),
        ))
        
        # Contact info style
        self.styles.add(ParagraphStyle(
            name='ResumeContactInfo',
            fontName='Helvetica',
            fontSize=10,
            leading=10,
            alignment=TA_CENTER,
            spaceAfter=6,
        ))
        
        # Bullet style
        self.styles.add(ParagraphStyle(
            name='ResumeBullet',
            fontName='Helvetica',
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
            leftIndent=15,
            spaceAfter=2,
        ))
    
    def create_pdf(self, resume_text, output_path):
        """
        Generate PDF from plain text resume.
        
        Args:
            resume_text: Plain text resume content
            output_path: Output PDF file path
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=0.5 * inch,
            leftMargin=0.5 * inch,
            topMargin=0.5 * inch,
            bottomMargin=0.5 * inch,
        )
        
        elements = []
        lines = resume_text.strip().split('\n')
        
        for line in lines:
            line = line.rstrip()
            
            if not line.strip():
                # Blank line - add small spacer
                elements.append(Spacer(1, 0.05 * inch))
            
            elif line.isupper() or line.endswith(':'):
                # Section header
                elements.append(Paragraph(line, self.styles['ResumeSectionHeader']))
            
            elif line.startswith('  ') or line.startswith('\t'):
                # Bullet point
                clean_line = line.lstrip()
                if clean_line.startswith('- ') or clean_line.startswith('• '):
                    clean_line = clean_line[2:].strip()
                elements.append(Paragraph('• ' + clean_line, self.styles['ResumeBullet']))
            
            else:
                # Regular text
                elements.append(Paragraph(line, self.styles['ResumeNormal']))
        
        try:
            doc.build(elements)
            logger.info(f"Saved PDF: {output_path}")
            print(f"Saved resume to: {output_path}")
        except Exception as e:
            logger.error(f"Failed to generate PDF {output_path}: {e}")
            raise


# ============================================================================
# File Manager - Organize resumes by date and company
# ============================================================================

class FileManager:
    """Manage file organization and naming"""

    INVALID_FILENAME_CHARS = re.compile(r'[<>:"/\\|?*]+')
    MULTISPACE_PATTERN = re.compile(r"\s+")
    MULTIDASH_PATTERN = re.compile(r"-{2,}")
    MAX_FILENAME_LENGTH = 120

    def __init__(self, base_path=None):
        """
        Initialize file manager.
        
        Args:
            base_path: Base directory for storing resumes
        """
        self.base_path = Path(base_path) if base_path else Path(config.RESUME_OUTPUT_PATH)
        self.base_path.mkdir(parents=True, exist_ok=True)
        logger.info(f"FileManager base path: {self.base_path}")

    def get_path(self, company, job_title, format='docx', date_str=None, source='Unknown'):
        """
        Get a versioned file path for a tailored resume.
        
        Args:
            company: Company name
            job_title: Job title
            format: 'docx', 'pdf', or 'json'
            date_str: Optional date string in YYYY-MM-DD format
            source: Job board/source label
        
        Returns:
            Path object for the file
        """
        date_part = self._normalize_date(date_str)
        base_name = self._build_base_name(
            date_part=date_part,
            company=company,
            role_title=job_title,
            source_job_board=source,
        )
        version = self._resolve_version(base_name=base_name, requested_extension=format)
        return self.base_path / f"{base_name}__v{version}.{format}"

    def get_metadata_path(self, company, job_title, date_str=None, source='Unknown'):
        """Get the sidecar JSON metadata path for a tailored resume."""
        return self.get_path(
            company=company,
            job_title=job_title,
            format='json',
            date_str=date_str,
            source=source,
        )

    def write_metadata(
        self,
        output_path,
        company,
        role_title,
        source_job_board,
        resume_type,
        source_url=None,
        created_at=None,
    ):
        """
        Write a sidecar JSON metadata file beside the generated resume.
        
        Returns:
            Path to the JSON metadata file
        """
        output_path = Path(output_path)
        metadata_path = output_path.with_suffix('.json')
        metadata = {
            'created_at': self._normalize_created_at(created_at),
            'company': company,
            'role_title': role_title,
            'source_job_board': source_job_board,
            'source_url': source_url,
            'resume_type': resume_type,
            'filename': output_path.name,
            'full_output_path': str(output_path.resolve()),
            'artifacts': self._collect_artifacts(output_path),
        }

        metadata_path.write_text(json.dumps(metadata, indent=2), encoding='utf-8')
        logger.info(f"Saved metadata: {metadata_path}")
        return metadata_path

    def save_docx_resume(
        self,
        docx_generator,
        company,
        job_title,
        source='Unknown',
        source_url=None,
        created_at=None,
    ):
        """Save a DOCX resume and create its metadata sidecar."""
        output_path = self.get_path(
            company=company,
            job_title=job_title,
            format='docx',
            date_str=created_at,
            source=source,
        )
        docx_generator.save(output_path)
        self.write_metadata(
            output_path=output_path,
            company=company,
            role_title=job_title,
            source_job_board=source,
            source_url=source_url,
            resume_type='docx',
            created_at=created_at,
        )
        return output_path

    def save_pdf_resume(
        self,
        pdf_generator,
        resume_text,
        company,
        job_title,
        source='Unknown',
        source_url=None,
        created_at=None,
    ):
        """Save a PDF resume and create its metadata sidecar."""
        output_path = self.get_path(
            company=company,
            job_title=job_title,
            format='pdf',
            date_str=created_at,
            source=source,
        )
        pdf_generator.create_pdf(resume_text, output_path)
        self.write_metadata(
            output_path=output_path,
            company=company,
            role_title=job_title,
            source_job_board=source,
            source_url=source_url,
            resume_type='pdf',
            created_at=created_at,
        )
        return output_path

    def _build_base_name(self, date_part, company, role_title, source_job_board):
        """Build a sortable filename stem."""
        company_part = self._sanitize_name(company, fallback='Company')
        role_part = self._sanitize_name(role_title, fallback='Role')
        source_part = self._sanitize_name(source_job_board, fallback='Unknown')

        version_suffix_length = len('__v999')
        available_length = self.MAX_FILENAME_LENGTH - version_suffix_length
        segments = [date_part, company_part, role_part, source_part]
        base_name = '__'.join(segments)

        while len(base_name) > available_length:
            trim_index = max(range(1, len(segments)), key=lambda idx: len(segments[idx]))
            minimum_lengths = {1: 8, 2: 12, 3: 6}
            if len(segments[trim_index]) <= minimum_lengths[trim_index]:
                break
            segments[trim_index] = segments[trim_index][:-1].rstrip('-')
            base_name = '__'.join(segments)

        return base_name[:available_length].rstrip('-.')

    def _resolve_version(self, base_name, requested_extension):
        """
        Find the next version number without splitting docx/pdf/json siblings
        across different versions.
        """
        version_map = {}
        pattern = re.compile(rf"^{re.escape(base_name)}__v(\d+)\.(docx|pdf|json)$", re.IGNORECASE)
        for candidate in self.base_path.glob(f"{base_name}__v*.*"):
            match = pattern.match(candidate.name)
            if not match:
                continue
            version = int(match.group(1))
            version_map.setdefault(version, set()).add(match.group(2).lower())

        if not version_map:
            return 1

        latest_version = max(version_map)
        existing_extensions = version_map[latest_version]
        if requested_extension.lower() not in existing_extensions:
            return latest_version

        return latest_version + 1

    def _collect_artifacts(self, output_path):
        """Collect sibling artifact paths that share the same base name."""
        artifact_paths = {}
        for extension in ('docx', 'pdf', 'json'):
            candidate = output_path.with_suffix(f'.{extension}')
            if candidate.exists():
                artifact_paths[extension] = str(candidate.resolve())
        return artifact_paths

    @staticmethod
    def _normalize_created_at(created_at):
        """Return an ISO timestamp string for metadata."""
        if isinstance(created_at, datetime):
            return created_at.isoformat(timespec='seconds')
        if isinstance(created_at, str) and created_at.strip():
            return created_at.strip()
        return datetime.now().isoformat(timespec='seconds')

    @staticmethod
    def _normalize_date(date_value=None):
        """Return the filename date segment in YYYY-MM-DD format."""
        if isinstance(date_value, datetime):
            return date_value.strftime('%Y-%m-%d')
        if isinstance(date_value, str) and date_value.strip():
            return date_value.strip()
        return datetime.now().strftime('%Y-%m-%d')

    @staticmethod
    def _sanitize_name(name, fallback='Unknown'):
        """Sanitize a filename segment for Windows-friendly output."""
        if name is None:
            return fallback

        clean = str(name).strip()
        clean = FileManager.INVALID_FILENAME_CHARS.sub(' ', clean)
        clean = FileManager.MULTISPACE_PATTERN.sub('-', clean)
        clean = FileManager.MULTIDASH_PATTERN.sub('-', clean)
        clean = clean.strip('-. _')
        return clean[:50] or fallback


# ============================================================================
# Example Usage
# ============================================================================

if __name__ == '__main__':
    from loguru import logger
    
    # Configure logging
    logger.remove()
    logger.add(lambda msg: print(msg, end=''), format="{message}", level="INFO")
    
    sample_resume = """
John Smith
john.smith@email.com | 555-123-4567 | LinkedIn.com/in/johnsmith

PROFESSIONAL SUMMARY
Experienced software engineer with 7+ years building scalable applications.

EXPERIENCE

Senior Software Engineer
Tech Company Inc., San Francisco, CA | Jan 2022 - Present
  - Led microservices platform handling 1M+ daily requests
  - Improved API response time by 60%
  - Mentored team of 4 junior engineers

EDUCATION

Bachelor of Science in Computer Science
University of California, Berkeley, CA | 2016

SKILLS
Languages: Python, JavaScript, SQL
Databases: PostgreSQL, MongoDB
Cloud: AWS, Docker
"""
    
    print("Testing DOCX Generator...")
    docx_gen = DocxGenerator()
    docx_gen.parse_and_add_resume(sample_resume)
    
    print("\nTesting PDF Generator...")
    pdf_gen = PdfGenerator()
    
    print("\nTesting File Manager...")
    fm = FileManager(config.RESUME_OUTPUT_PATH)
    docx_path = fm.save_docx_resume(
        docx_generator=docx_gen,
        company='Tech Corp',
        job_title='Senior Engineer',
        source='Demo',
        source_url='https://example.com/jobs/1',
        created_at='2026-03-27',
    )
    pdf_path = fm.save_pdf_resume(
        pdf_generator=pdf_gen,
        resume_text=sample_resume,
        company='Tech Corp',
        job_title='Senior Engineer',
        source='Demo',
        source_url='https://example.com/jobs/1',
        created_at='2026-03-27',
    )
    print(f"Generated DOCX path: {docx_path}")
    print(f"Generated PDF path: {pdf_path}")
    print(f"Generated metadata path: {docx_path.with_suffix('.json')}")
